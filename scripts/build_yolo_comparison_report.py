from __future__ import annotations

import argparse
import csv
from pathlib import Path
from typing import Any

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build a DOCX report from YOLO comparison metrics.")
    parser.add_argument("--metrics", default=Path("runs/compare_eval/yolo_model_comparison.csv"), type=Path)
    parser.add_argument("--output", default=Path("Digital AI Proposed Project - YOLO Comparison.docx"), type=Path)
    parser.add_argument("--figures-dir", default=Path("runs/compare_eval/figures"), type=Path)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    rows = read_metrics(args.metrics)
    if not rows:
        raise SystemExit(f"No rows found in {args.metrics}")

    args.figures_dir.mkdir(parents=True, exist_ok=True)
    map_figure = plot_map_comparison(rows, args.figures_dir)
    resource_figure = plot_resource_tradeoff(rows, args.figures_dir)
    latency_figure = plot_latency(rows, args.figures_dir)

    selected = select_edge_model(rows)
    best_external = max(rows, key=lambda row: row["cvc_map50_95"])

    doc = Document()
    make_landscape(doc)
    set_base_style(doc)

    title = doc.add_heading("YOLO Model Comparison for Real-Time Polyp Detection", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Supplement to the proposed gastrointestinal AI final project. The comparison uses the "
        f"following trained candidates: {format_model_list(rows)}."
    )

    doc.add_heading("Decision Summary", level=1)
    doc.add_paragraph(
        f"The recommended deployment candidate is {selected['model']}. It gives the best balance "
        f"between external-test performance and edge feasibility: {selected['params_m']:.2f}M "
        f"parameters, {selected['gflops']:.2f} GFLOPs, CVC-ClinicDB mAP50-95 "
        f"{selected['cvc_map50_95']:.3f}, and ONNX Runtime CPU latency "
        f"{selected['onnx_cpu_mean_ms']:.2f} ms per 416 x 416 image on this workstation."
    )
    doc.add_paragraph(
        f"The highest external-test mAP50-95 in this run is {best_external['model']} "
        f"({best_external['cvc_map50_95']:.3f}). The final choice favors a compact model unless "
        "the larger model provides a clinically meaningful external-validation gain that offsets "
        "its additional memory and compute cost."
    )

    doc.add_heading("Experimental Setup", level=1)
    setup = doc.add_table(rows=1, cols=2)
    setup.style = "Table Grid"
    setup.rows[0].cells[0].text = "Item"
    setup.rows[0].cells[1].text = "Setting"
    add_kv_row(setup, "Training dataset", "Kvasir-SEG training split, converted from masks to YOLO bounding boxes")
    add_kv_row(setup, "Internal validation", "Kvasir-SEG validation split")
    add_kv_row(setup, "External validation", "CVC-ClinicDB test split")
    add_kv_row(setup, "Image size", "416 x 416")
    add_kv_row(setup, "Training schedule", "60 epochs, batch size 32, same preprocessing and augmentation settings")
    add_kv_row(setup, "Latency measurement", "Exported ONNX model, ONNX Runtime CPUExecutionProvider, inference only")

    doc.add_heading("Main Results", level=1)
    add_results_table(doc, rows)

    doc.add_heading("Figures", level=1)
    doc.add_picture(str(map_figure), width=Inches(8.2))
    doc.add_paragraph("Figure 1. Internal and external mAP50-95 comparison.")
    doc.add_picture(str(resource_figure), width=Inches(8.2))
    doc.add_paragraph("Figure 2. Parameter count versus external-test mAP50-95.")
    doc.add_picture(str(latency_figure), width=Inches(8.2))
    doc.add_paragraph("Figure 3. ONNX Runtime CPU latency and throughput comparison.")

    doc.add_heading("Interpretation", level=1)
    doc.add_paragraph(
        "The external CVC-ClinicDB test is the most important result for classroom justification "
        "because it tests whether the detector generalizes beyond the Kvasir-SEG training source. "
        "For an STM32MP257-oriented project, parameter count, ONNX size, and latency are also part "
        "of the selection criteria because the final target is an embedded NPU workflow."
    )
    doc.add_paragraph(
        f"Based on this run, {selected['model']} should be used as the primary edge-deployment "
        f"candidate. {format_alternative_sentence(rows, selected)}"
    )
    doc.add_paragraph(
        "The workstation ONNX latency is not a substitute for STM32MP257 NPU latency. It is used "
        "here as a reproducible relative benchmark. The next implementation step is to convert the "
        "selected ONNX model through the STM32 deployment toolchain and measure real NPU throughput "
        "on endoscopy-like video frames."
    )

    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)
    print(f"Wrote {args.output}")


def read_metrics(path: Path) -> list[dict[str, Any]]:
    with path.open("r", encoding="utf-8", newline="") as handle:
        reader = csv.DictReader(handle)
        rows = []
        for row in reader:
            parsed: dict[str, Any] = {}
            for key, value in row.items():
                parsed[key] = parse_value(value)
            rows.append(parsed)
        return rows


def parse_value(value: str) -> Any:
    try:
        return float(value)
    except (TypeError, ValueError):
        return value


def make_landscape(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)


def set_base_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(9)


def add_kv_row(table: Any, key: str, value: str) -> None:
    cells = table.add_row().cells
    cells[0].text = key
    cells[1].text = value


def add_results_table(doc: Document, rows: list[dict[str, Any]]) -> None:
    headers = [
        "Model",
        "Params (M)",
        "GFLOPs",
        "PT MB",
        "ONNX MB",
        "Kvasir P/R",
        "Kvasir mAP50",
        "Kvasir mAP50-95",
        "CVC P/R",
        "CVC mAP50",
        "CVC mAP50-95",
        "CPU ms",
        "FPS",
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
    for row in rows:
        values = [
            row["model"],
            f"{row['params_m']:.2f}",
            f"{row['gflops']:.2f}",
            f"{row['pt_size_mb']:.2f}",
            f"{row['onnx_size_mb']:.2f}",
            f"{row['kvasir_precision']:.3f}/{row['kvasir_recall']:.3f}",
            f"{row['kvasir_map50']:.3f}",
            f"{row['kvasir_map50_95']:.3f}",
            f"{row['cvc_precision']:.3f}/{row['cvc_recall']:.3f}",
            f"{row['cvc_map50']:.3f}",
            f"{row['cvc_map50_95']:.3f}",
            f"{row['onnx_cpu_mean_ms']:.2f}",
            f"{row['onnx_cpu_fps']:.1f}",
        ]
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = value


def format_model_list(rows: list[dict[str, Any]]) -> str:
    names = [str(row["model"]) for row in rows]
    if len(names) == 1:
        return names[0]
    if len(names) == 2:
        return f"{names[0]} and {names[1]}"
    return f"{', '.join(names[:-1])}, and {names[-1]}"


def format_alternative_sentence(rows: list[dict[str, Any]], selected: dict[str, Any]) -> str:
    alternatives = [str(row["model"]) for row in rows if row["model"] != selected["model"]]
    if not alternatives:
        return "No alternative model was included in this metrics file."
    return (
        f"The remaining evaluated model(s), {', '.join(alternatives)}, should be kept as ablation "
        "baselines to show how clinical performance changes with model family, augmentation, and "
        "compute budget."
    )


def select_edge_model(rows: list[dict[str, Any]]) -> dict[str, Any]:
    compact = [row for row in rows if row["params_m"] <= 3.5]
    candidates = compact if compact else rows
    return max(candidates, key=lambda row: (row["cvc_map50_95"], row["cvc_recall"], -row["onnx_cpu_mean_ms"]))


def plot_map_comparison(rows: list[dict[str, Any]], figures_dir: Path) -> Path:
    labels = [row["model"] for row in rows]
    x = np_arange(len(labels))
    width = 0.36
    fig, ax = plt.subplots(figsize=(8.4, 4.2))
    ax.bar([value - width / 2 for value in x], [row["kvasir_map50_95"] for row in rows], width, label="Kvasir val")
    ax.bar([value + width / 2 for value in x], [row["cvc_map50_95"] for row in rows], width, label="CVC external")
    ax.set_ylabel("mAP50-95")
    ax.set_ylim(0, 1.0)
    ax.set_xticks(x)
    ax.set_xticklabels(labels)
    ax.legend()
    ax.grid(axis="y", alpha=0.25)
    fig.tight_layout()
    path = figures_dir / "map50_95_comparison.png"
    fig.savefig(path, dpi=220)
    plt.close(fig)
    return path


def plot_resource_tradeoff(rows: list[dict[str, Any]], figures_dir: Path) -> Path:
    fig, ax = plt.subplots(figsize=(8.4, 4.2))
    x = [row["params_m"] for row in rows]
    y = [row["cvc_map50_95"] for row in rows]
    sizes = [max(80, row["onnx_cpu_mean_ms"] * 10) for row in rows]
    ax.scatter(x, y, s=sizes)
    for row in rows:
        ax.annotate(row["model"], (row["params_m"], row["cvc_map50_95"]), xytext=(5, 5), textcoords="offset points")
    ax.set_xlabel("Parameters (M)")
    ax.set_ylabel("CVC external mAP50-95")
    ax.set_ylim(0, 1.0)
    ax.grid(alpha=0.25)
    fig.tight_layout()
    path = figures_dir / "resource_tradeoff.png"
    fig.savefig(path, dpi=220)
    plt.close(fig)
    return path


def plot_latency(rows: list[dict[str, Any]], figures_dir: Path) -> Path:
    labels = [row["model"] for row in rows]
    x = np_arange(len(labels))
    fig, ax1 = plt.subplots(figsize=(8.4, 4.2))
    ax1.bar(x, [row["onnx_cpu_mean_ms"] for row in rows], label="Mean latency (ms)")
    ax1.set_ylabel("Mean latency (ms)")
    ax1.set_xticks(x)
    ax1.set_xticklabels(labels)
    ax1.grid(axis="y", alpha=0.25)

    ax2 = ax1.twinx()
    ax2.plot(x, [row["onnx_cpu_fps"] for row in rows], color="black", marker="o", label="FPS")
    ax2.set_ylabel("FPS")
    fig.tight_layout()
    path = figures_dir / "latency_fps.png"
    fig.savefig(path, dpi=220)
    plt.close(fig)
    return path


def np_arange(length: int) -> list[float]:
    return [float(index) for index in range(length)]


if __name__ == "__main__":
    main()
