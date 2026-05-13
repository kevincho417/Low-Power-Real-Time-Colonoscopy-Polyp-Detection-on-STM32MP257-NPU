from __future__ import annotations

import csv
import shutil
from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


PROJECT = Path.cwd()
REPORT = PROJECT / "Digital AI Proposed Project - YOLO Comparison.docx"
BACKUP = PROJECT / "Digital AI Proposed Project - YOLO Comparison.previous-results.bak.docx"
SUMMARY_CSV = PROJECT / "runs/multidata_compare_eval_4models/model_summary.csv"
DETAIL_CSV = PROJECT / "runs/multidata_compare_eval_4models/multidata_eval.csv"
FIG_DIR = PROJECT / "runs/multidata_compare_eval_4models/report_figures"


def main() -> None:
    summary = read_csv(SUMMARY_CSV)
    details = read_csv(DETAIL_CSV)
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    figures = make_figures(summary)

    if REPORT.exists() and not BACKUP.exists():
        shutil.copy2(REPORT, BACKUP)

    doc = Document()
    add_title(doc)
    add_decision_summary(doc)
    add_experimental_setup(doc)
    add_main_table(doc, summary)
    add_external_detail_table(doc, details)
    add_figures(doc, figures)
    add_interpretation(doc)
    add_files_section(doc)
    doc.save(REPORT)
    print(f"Updated {REPORT}")
    print(f"Backup {BACKUP}")


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8", newline="") as handle:
        return list(csv.DictReader(handle))


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("YOLO Model Comparison for Real-Time Polyp Detection")
    run.bold = True
    run.font.size = Pt(16)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        "Updated multi-dataset comparison for the gastrointestinal AI final project "
        "(YOLO26n, YOLO11n, YOLOv8n, and YOLOv8s)."
    )


def add_decision_summary(doc: Document) -> None:
    doc.add_heading("Decision Summary", level=1)
    doc.add_paragraph(
        "The current multi-dataset experiment changes the interpretation of the previous "
        "Kvasir-only comparison. YOLOv8s obtains the best cross-dataset accuracy, with an "
        "external average mAP50-95 of 0.502, but it is substantially larger at 11.14M "
        "parameters. Among nano-scale models, YOLO11n has the best external average "
        "mAP50-95 at 0.481, narrowly ahead of YOLO26n and YOLOv8n."
    )
    doc.add_paragraph(
        "For STM32MP257 deployment, YOLOv8n remains the primary demonstration model because "
        "its full-integer INT8 TFLite-to-NPU path has already been validated on the board. "
        "It runs on the STM32MP257 NPU at approximately 35.7 ms per inference, around 28 FPS. "
        "YOLOv8s is kept as the accuracy upper-bound reference, while YOLO11n and YOLO26n "
        "are compact alternatives for additional conversion tests."
    )


def add_experimental_setup(doc: Document) -> None:
    doc.add_heading("Experimental Setup", level=1)
    rows = [
        ("Training data", "Multi-dataset training set: Kvasir-SEG, PolypGen positive cropped images, Hyper-Kvasir segmented positives, and Hyper-Kvasir lower-GI negatives."),
        ("Internal holdout", "Held-out test split from the multi-dataset configuration."),
        ("External validation", "CVC-ClinicDB, CVC-ColonDB, and ETIS test sets."),
        ("Task", "Single-class polyp bounding-box detection."),
        ("Image size", "416 x 416 for all reported accuracy experiments."),
        ("Training schedule", "80 epochs, batch size 32, same endoscopy-domain augmentation and seed where applicable."),
        ("Metrics", "Precision, recall, mAP50, mAP50-95, parameters, GFLOPs, and model weight size."),
        ("Deployment context", "YOLOv8n full-integer INT8 is the currently validated STM32MP257 NPU deployment model."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Item"
    table.rows[0].cells[1].text = "Setting"
    for item, setting in rows:
        cells = table.add_row().cells
        cells[0].text = item
        cells[1].text = setting


def add_main_table(doc: Document, summary: list[dict[str, str]]) -> None:
    doc.add_heading("Main Results", level=1)
    headers = [
        "Model",
        "Params (M)",
        "GFLOPs",
        "PT MB",
        "Internal mAP50-95",
        "External Avg mAP50-95",
        "External Avg Recall",
        "CVC mAP50-95",
        "ColonDB mAP50-95",
        "ETIS mAP50-95",
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
    for row in summary:
        cells = table.add_row().cells
        values = [
            row["model"].replace("-MultiData", ""),
            f"{float(row['params_m']):.2f}",
            f"{float(row['gflops']):.2f}",
            f"{float(row['weight_mb']):.2f}",
            f"{float(row['internal_map5095']):.3f}",
            f"{float(row['external_avg_map5095']):.3f}",
            f"{float(row['external_avg_recall']):.3f}",
            f"{float(row['clinicdb_map5095']):.3f}",
            f"{float(row['colondb_map5095']):.3f}",
            f"{float(row['etis_map5095']):.3f}",
        ]
        for cell, value in zip(cells, values):
            cell.text = value


def add_external_detail_table(doc: Document, details: list[dict[str, str]]) -> None:
    doc.add_heading("Per-Dataset Accuracy Details", level=1)
    headers = ["Model", "Dataset", "Precision", "Recall", "mAP50", "mAP50-95"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
    for row in details:
        cells = table.add_row().cells
        values = [
            row["model"].replace("-MultiData", ""),
            row["dataset"],
            f"{float(row['precision']):.3f}",
            f"{float(row['recall']):.3f}",
            f"{float(row['map50']):.3f}",
            f"{float(row['map50_95']):.3f}",
        ]
        for cell, value in zip(cells, values):
            cell.text = value


def make_figures(summary: list[dict[str, str]]) -> list[tuple[Path, str]]:
    labels = [row["model"].replace("-MultiData", "") for row in summary]
    internal = [float(row["internal_map5095"]) for row in summary]
    external = [float(row["external_avg_map5095"]) for row in summary]
    params = [float(row["params_m"]) for row in summary]

    fig1 = FIG_DIR / "map5095_internal_external.png"
    x = range(len(labels))
    plt.figure(figsize=(8, 4.5))
    plt.bar([i - 0.18 for i in x], internal, width=0.36, label="Internal holdout")
    plt.bar([i + 0.18 for i in x], external, width=0.36, label="External average")
    plt.xticks(list(x), labels, rotation=15, ha="right")
    plt.ylabel("mAP50-95")
    plt.ylim(0, 0.85)
    plt.legend()
    plt.tight_layout()
    plt.savefig(fig1, dpi=220)
    plt.close()

    fig2 = FIG_DIR / "params_vs_external_map.png"
    plt.figure(figsize=(7, 4.5))
    plt.scatter(params, external, s=90)
    for label, p, m in zip(labels, params, external):
        plt.annotate(label, (p, m), textcoords="offset points", xytext=(6, 5))
    plt.xlabel("Parameters (M)")
    plt.ylabel("External average mAP50-95")
    plt.grid(True, alpha=0.3)
    plt.tight_layout()
    plt.savefig(fig2, dpi=220)
    plt.close()

    fig3 = FIG_DIR / "external_dataset_map5095.png"
    clinic = [float(row["clinicdb_map5095"]) for row in summary]
    colon = [float(row["colondb_map5095"]) for row in summary]
    etis = [float(row["etis_map5095"]) for row in summary]
    plt.figure(figsize=(8, 4.5))
    plt.bar([i - 0.25 for i in x], clinic, width=0.25, label="CVC-ClinicDB")
    plt.bar(list(x), colon, width=0.25, label="CVC-ColonDB")
    plt.bar([i + 0.25 for i in x], etis, width=0.25, label="ETIS")
    plt.xticks(list(x), labels, rotation=15, ha="right")
    plt.ylabel("mAP50-95")
    plt.ylim(0, 0.65)
    plt.legend()
    plt.tight_layout()
    plt.savefig(fig3, dpi=220)
    plt.close()

    return [
        (fig1, "Figure 1. Internal holdout and external-average mAP50-95 under the current multi-dataset training setup."),
        (fig2, "Figure 2. Parameter count versus external-average mAP50-95. YOLOv8s improves accuracy but at a much larger parameter cost."),
        (fig3, "Figure 3. External dataset mAP50-95 comparison on CVC-ClinicDB, CVC-ColonDB, and ETIS."),
    ]


def add_figures(doc: Document, figures: list[tuple[Path, str]]) -> None:
    doc.add_heading("Figures", level=1)
    for path, caption in figures:
        doc.add_picture(str(path), width=Inches(6.3))
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_interpretation(doc: Document) -> None:
    doc.add_heading("Interpretation", level=1)
    doc.add_paragraph(
        "The external datasets are the most important evidence for classroom justification "
        "because they test cross-dataset generalization rather than only within-distribution "
        "performance. The current results show that multi-dataset training improves the "
        "scientific strength of the project: the model is evaluated on multiple independent "
        "polyp datasets with different image characteristics."
    )
    doc.add_paragraph(
        "YOLOv8s is the best pure-accuracy model in this run, but it is not the preferred "
        "embedded deployment candidate because its parameter count and compute are much higher. "
        "YOLOv8n is selected for the live STM32MP257 demonstration because its full-integer INT8 "
        "model has already been converted to NPU format and benchmarked successfully. YOLO11n "
        "is a promising compact accuracy candidate for future NPU conversion, while YOLO26n is "
        "compact but its tested INT8 NPU model was slower and less stable in video inference."
    )


def add_files_section(doc: Document) -> None:
    doc.add_heading("Generated Result Files", level=1)
    for text in [
        "runs/multidata_compare_eval_4models/model_summary.csv",
        "runs/multidata_compare_eval_4models/multidata_eval.csv",
        "runs/multidata_compare_eval_4models/multidata_eval.json",
        "runs/detect/runs/multidata_train/multidata_yolo11n_416/weights/best.pt",
        "runs/detect/runs/multidata_train/multidata_yolo26n_416/weights/best.pt",
        "runs/detect/runs/multidata_train/multidata_yolov8n_416/weights/best.pt",
        "runs/detect/runs/multidata_train/multidata_yolov8s_416/weights/best.pt",
    ]:
        doc.add_paragraph(text, style="List Bullet")


if __name__ == "__main__":
    main()
