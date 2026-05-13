from __future__ import annotations

import csv
from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


PROJECT = Path.cwd()
OUT_DOCX = PROJECT / "Polyp Detection STM32MP257 Full Paper Report.docx"
RESULT_DIR = PROJECT / "runs/multidata_compare_eval_4models"
SUMMARY_CSV = RESULT_DIR / "model_summary.csv"
DETAIL_CSV = RESULT_DIR / "multidata_eval.csv"
FIG_DIR = RESULT_DIR / "full_paper_figures"
DOC_IMAGE_DIR = PROJECT / "DOC images"


DATASET_ROWS = [
    ("Kvasir-SEG", "Image-level colonoscopy polyp dataset with segmentation masks", "Converted to YOLO bounding boxes; used in multi-dataset training/validation"),
    ("PolypGen positive cropped", "Positive polyp images with regional annotations", "Used as additional positive-domain training and test data"),
    ("Hyper-Kvasir segmented positives", "Lower-GI polyp images with segmentation annotations", "Used as additional positive-domain training and test data"),
    ("Hyper-Kvasir lower-GI negatives", "Lower-GI non-polyp frames", "Used as negative/background examples to reduce false positives"),
    ("CVC-ClinicDB", "612 colonoscopy frames with polyp masks", "External validation dataset"),
    ("CVC-ColonDB", "380 colonoscopy frames with polyp masks", "External validation dataset"),
    ("ETIS-LaribPolypDB", "196 colonoscopy frames with polyp masks", "External validation dataset"),
]


TRAINING_ROWS = [
    ("Input size", "416 x 416 for the main accuracy comparison"),
    ("Models", "YOLO11n, YOLO26n, YOLOv8n, YOLOv8s"),
    ("Epochs", "80"),
    ("Batch size", "32"),
    ("Optimizer", "Ultralytics auto optimizer selection"),
    ("Seed", "42"),
    ("Augmentation", "HSV jitter, translation, scale jitter, shear, perspective, horizontal/vertical flip, mosaic, mixup"),
    ("Task", "Single-class polyp detection"),
    ("Evaluation", "Internal holdout plus external CVC-ClinicDB, CVC-ColonDB, and ETIS"),
]


DEPLOYMENT_ROWS = [
    ("YOLOv8n ONNX CPU on STM32MP257", "FP32 ONNX Runtime CPU", "879.63 ms/frame", "1.14 FPS", "Baseline fallback; not real-time"),
    ("YOLO26n previous .nb", "NPU, float16 I/O", "422.70 ms", "2.37 FPS", "Not true INT8 behavior; rejected"),
    ("YOLO11n full-integer .nb", "NPU, int8 I/O", "43.20 ms", "23.15 FPS", "NPU validated; detection stability close to YOLOv8n but slower"),
    ("YOLO26n full-integer .nb", "NPU, int8 I/O", "44.08 ms", "22.69 FPS", "Runs on NPU but lower video-detection stability"),
    ("YOLOv8n full-integer .nb", "NPU, int8 I/O", "35.70 ms", "28.01 FPS", "Selected embedded demo model"),
]


PIPELINE_ROWS = [
    ("Original Python pipeline", "Standard resize, RGB conversion, float normalization, quantization, NMS", "51.64", "36.36", "1.47", "10.68", "87/200"),
    ("Fast INT8 letterbox pipeline", "Direct INT8 tensor creation while preserving YOLO letterbox geometry", "17.47", "36.59", "1.64", "17.44", "87/200"),
    ("GStreamer direct resize", "GStreamer decode/resize before Python; direct 416x416 stretch", "7.60", "37.48", "1.20", "21.19", "45/200"),
]


QUANTIZATION_ROWS = [
    ("YOLOv8n 416 full-integer TFLite", "Input int8 [1,416,416,3]; output int8 [1,5,3549]", "Converted successfully to .nb; NPU validated"),
    ("YOLOv8n 384 full-integer TFLite", "Input int8 [1,384,384,3]; output int8 [1,5,3024]", "Prepared for future speed/accuracy trade-off testing"),
    ("YOLOv8n 320 full-integer TFLite", "Input int8 [1,320,320,3]; output int8 [1,5,2100]", "Prepared for aggressive speed testing"),
    ("YOLO11n full-integer TFLite/.nb", "Input int8 [1,416,416,3]; output int8 [1,5,3549]", "Converted to .nb and validated on STM32MP257 NPU"),
    ("YOLO26n full-integer TFLite", "Input int8 [1,416,416,3]; output int8 [1,300,6]", "Converted to .nb and validated, but inferior to YOLOv8n in deployment tests"),
    ("ONNX QDQ INT8 variants", "Several QDQ and conv-only variants tested", "ST optimizer compatibility was unreliable; TFLite route was preferred"),
]


REFERENCES = [
    "World Health Organization. Colorectal cancer fact sheet. https://www.who.int/news-room/fact-sheets/detail/colorectal-cancer",
    "CDC. Screening for Colorectal Cancer. https://www.cdc.gov/colorectal-cancer/screening/",
    "Kvasir-SEG dataset. Simula. https://datasets.simula.no/kvasir-seg/",
    "CVC-ClinicDB dataset. https://polyp.grand-challenge.org/CVCClinicDB/",
    "STMicroelectronics. STM32MP257 product family and NPU documentation. https://www.st.com/",
    "Ultralytics YOLO documentation. https://docs.ultralytics.com/",
]


def main() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    summary = read_csv(SUMMARY_CSV)
    detail = read_csv(DETAIL_CSV)
    figures = make_figures(summary)

    doc = Document()
    set_margins(doc)
    add_title(doc)
    add_abstract(doc)
    add_keywords(doc)
    add_introduction(doc)
    add_materials_methods(doc)
    add_results(doc, summary, detail, figures)
    add_discussion(doc)
    add_limitations(doc)
    add_conclusion(doc)
    add_reproducibility(doc)
    add_references(doc)
    doc.save(OUT_DOCX)
    print(f"wrote {OUT_DOCX}")


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8", newline="") as handle:
        return list(csv.DictReader(handle))


def set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Low-Power Real-Time Colonoscopy Polyp Detection on STM32MP257 NPU")
    r.bold = True
    r.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("A Multi-Dataset YOLO Training, Quantization, and Edge Deployment Study")
    r.italic = True
    r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Biomedical AI Final Project Report | May 2026")


def add_abstract(doc: Document) -> None:
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "Colonoscopy polyp detection is a clinically meaningful computer-aided detection "
        "task because the output can directly guide endoscopist attention during a procedure. "
        "This project developed a lightweight edge-AI prototype for real-time polyp detection "
        "using public gastrointestinal datasets and STM32MP257 NPU deployment. Segmentation "
        "annotations from multiple datasets were converted into YOLO bounding boxes, and four "
        "YOLO variants were trained under the same multi-dataset protocol: YOLO11n, YOLO26n, "
        "YOLOv8n, and YOLOv8s. Cross-dataset evaluation was performed on internal holdout data, "
        "CVC-ClinicDB, CVC-ColonDB, and ETIS. YOLOv8s achieved the best external average "
        "mAP50-95 of 0.502, while YOLO11n was the strongest nano-scale accuracy candidate "
        "with external average mAP50-95 of 0.481. For deployment, YOLOv8n was selected because "
        "its full-integer INT8 TFLite model was successfully converted to STM32MP257 NPU format "
        "and benchmarked at approximately 35.7 ms per inference, corresponding to 28.0 FPS. "
        "The final prototype demonstrates a complete path from public medical datasets to "
        "quantized edge inference, while explicitly documenting accuracy, latency, and "
        "deployment trade-offs."
    )


def add_keywords(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run("Keywords: ").bold = True
    p.add_run("colonoscopy, polyp detection, YOLO, STM32MP257, NPU, INT8 quantization, edge AI, computer-aided detection")


def add_introduction(doc: Document) -> None:
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "Colorectal cancer is a major global health burden, and colonoscopy is clinically "
        "important because it enables both detection and removal of premalignant lesions. "
        "Computer-aided detection (CADe) systems are already clinically relevant in endoscopy, "
        "but most commercial systems are high-end and closed. A course-scale project can be "
        "clinically convincing if it studies the core workflow: detection of suspected polyps "
        "from live-like colonoscopy frames, quantitative external validation, and real-time "
        "deployment feasibility on a constrained edge device."
    )
    doc.add_paragraph(
        "The goal of this project was not to claim clinical readiness. Instead, the goal was "
        "to build a rigorous research prototype and measure whether a lightweight YOLO detector "
        "can preserve useful cross-dataset accuracy while running in real time on the "
        "STM32MP257 NPU after INT8 quantization."
    )
    add_optional_figure(
        doc,
        DOC_IMAGE_DIR / "Overall structure.png",
        "Figure A. Overall system architecture of the proposed colonoscopy polyp detection workflow.",
        "The overall architecture summarizes the complete project path from public endoscopy "
        "datasets to model training, quantization, STM32MP257 NPU deployment, and real-time "
        "video output. This figure is included early because it connects the clinical task, "
        "the machine-learning workflow, and the embedded-AI feasibility study in a single view.",
    )


def add_materials_methods(doc: Document) -> None:
    doc.add_heading("2. Materials and Methods", level=1)

    doc.add_heading("2.1 Datasets", level=2)
    doc.add_paragraph(
        "Multiple public gastrointestinal datasets were used to reduce dependence on a single "
        "source distribution. Segmentation masks were converted into YOLO bounding boxes by "
        "extracting the bounding rectangle around each positive mask. Negative lower-GI frames "
        "were included as empty-label images to improve background rejection."
    )
    add_table(doc, ["Dataset", "Material", "Role"], DATASET_ROWS)
    add_optional_figure(
        doc,
        DOC_IMAGE_DIR / "Data_composition&usage.png",
        "Figure B. Dataset composition and usage across training, validation, internal testing, and external testing.",
        "This figure explains why the datasets were not all used in the same way. Kvasir-SEG, "
        "PolypGen, Hyper-Kvasir segmented positives, and Hyper-Kvasir negative lower-GI frames "
        "formed the multi-dataset training distribution, while CVC-ClinicDB, CVC-ColonDB, and "
        "ETIS were preserved as external tests to assess cross-dataset generalization.",
    )

    doc.add_heading("2.2 Annotation Conversion and Data Format", level=2)
    doc.add_paragraph(
        "All training data were represented in YOLO detection format. Each image had a "
        "corresponding label file containing normalized class, center-x, center-y, width, and "
        "height values. Empty label files were retained for negative frames. The final "
        "multi-dataset configuration contained 6,127 training images, 775 validation images, "
        "and 657 internal test images, with both positive polyp frames and negative lower-GI "
        "background frames."
    )
    add_optional_figure(
        doc,
        DOC_IMAGE_DIR / "data_preprosessing_pipeline.png",
        "Figure C. Data preprocessing and annotation harmonization pipeline.",
        "The preprocessing pipeline standardizes heterogeneous public datasets into a single "
        "YOLO detection format. Segmentation masks are transformed into bounding boxes, "
        "image-label pairs are checked, negative frames are represented by empty label files, "
        "and the final data are organized into train, validation, internal-test, and external-test subsets.",
    )

    doc.add_heading("2.3 Model Training", level=2)
    add_table(doc, ["Item", "Setting"], TRAINING_ROWS)
    doc.add_paragraph(
        "The four candidate models were trained with identical image size and augmentation "
        "settings to make the comparison meaningful. YOLOv8s was included as an accuracy "
        "upper-bound model, while YOLO11n, YOLO26n, and YOLOv8n represented nano-scale "
        "deployment candidates."
    )
    add_optional_figure(
        doc,
        DOC_IMAGE_DIR / "Training & Quantize pipeline.png",
        "Figure D. Model training, quantization, and deployment conversion workflow.",
        "This workflow shows how the trained YOLO checkpoints were evaluated in PyTorch, "
        "exported to deployment formats, converted to full-integer INT8 TFLite when possible, "
        "and finally transformed into STM32MP257-compatible .nb files for NPU inference.",
    )

    doc.add_heading("2.4 Evaluation Protocol", level=2)
    doc.add_paragraph(
        "Accuracy was measured using precision, recall, mAP50, and mAP50-95. The internal "
        "holdout result estimates performance within the assembled multi-dataset distribution. "
        "External validation on CVC-ClinicDB, CVC-ColonDB, and ETIS was treated as the key "
        "generalization test because it evaluates domain shift across independent datasets."
    )

    doc.add_heading("2.5 Quantization and STM32MP257 Deployment", level=2)
    doc.add_paragraph(
        "Deployment experiments used TFLite full-integer INT8 quantization when possible. "
        "The TFLite route was more reliable than ONNX QDQ quantization for STM32MP257 NPU "
        "optimization. Models were converted into .nb files using ST Edge AI tooling and "
        "executed through the STAI MPU runtime on the STM32MP257 board."
    )
    add_table(doc, ["Artifact", "Tensor format", "Outcome"], QUANTIZATION_ROWS)


def add_results(doc: Document, summary: list[dict[str, str]], detail: list[dict[str, str]], figures: list[tuple[Path, str, str]]) -> None:
    doc.add_heading("3. Results", level=1)

    doc.add_heading("3.1 Cross-Dataset Accuracy", level=2)
    doc.add_paragraph(
        "YOLOv8s achieved the best external average mAP50-95, but with a much larger model. "
        "Among compact models, YOLO11n achieved the highest external average mAP50-95, while "
        "YOLOv8n remained the most reliable deployment candidate after NPU conversion."
    )
    add_summary_table(doc, summary)

    doc.add_heading("3.2 Per-Dataset Model Performance", level=2)
    add_detail_table(doc, detail)

    doc.add_heading("3.3 Figures", level=2)
    for path, caption, explanation in figures:
        add_optional_figure(doc, path, caption, explanation)

    add_model_selection_section(doc, summary)

    doc.add_heading("3.5 Embedded Inference and NPU Deployment", level=2)
    add_table(doc, ["Experiment", "Runtime format", "Latency", "FPS", "Interpretation"], DEPLOYMENT_ROWS)

    doc.add_heading("3.6 End-to-End Video Pipeline Optimization", level=2)
    doc.add_paragraph(
        "NPU inference alone was not the only bottleneck. Python preprocessing initially "
        "dominated the end-to-end video pipeline. A direct INT8 letterbox preprocessing path "
        "reduced preprocessing cost while preserving the same detection behavior. A GStreamer "
        "resize pipeline was faster but changed image geometry and reduced detections on the "
        "sample video, so it was not selected as the final demonstration pipeline."
    )
    add_table(doc, ["Pipeline", "Description", "Preprocess ms", "NPU ms", "Decode/NMS ms", "Wall FPS", "Detected frames"], PIPELINE_ROWS)
    add_optional_figure(
        doc,
        DOC_IMAGE_DIR / "STM32MP257_inference_pipeline.png",
        "Figure E. Two STM32MP257 inference pipelines evaluated for end-to-end video deployment.",
        "The pipeline diagram clarifies that NPU inference time is only one part of the "
        "embedded system. The final selected route keeps YOLO letterbox geometry and generates "
        "the INT8 tensor directly, while the faster GStreamer direct-resize route sacrifices "
        "geometric consistency and reduced detections in the sample video.",
    )


def add_model_selection_section(doc: Document, summary: list[dict[str, str]]) -> None:
    doc.add_heading("3.4 Model Selection and Accuracy-Latency Trade-off", level=2)
    doc.add_paragraph(
        "The final model was selected by combining two types of evidence rather than by using "
        "mAP alone. First, the parameter-versus-external-mAP figure shows the accuracy and "
        "model-complexity relationship. Second, the STM32MP257 deployment table measures the "
        "actual NPU latency and video-detection stability of the converted INT8 .nb models."
    )
    doc.add_paragraph(
        "YOLOv8s achieved the strongest accuracy result, with external average mAP50-95 of "
        "0.502 and internal mAP50-95 of 0.774, but it used 11.14M parameters and therefore "
        "served mainly as an accuracy upper bound. YOLO11n was the strongest nano-scale "
        "accuracy model, with external average mAP50-95 of 0.481 and 2.59M parameters, but "
        "its STM32MP257 NPU latency was 43.20 ms. YOLO26n was similarly compact but showed "
        "weaker deployment stability in the sample video. YOLOv8n had slightly lower external "
        "average mAP50-95 of 0.465, but it provided the fastest verified NPU inference at "
        "35.70 ms and the most defensible complete deployment path."
    )
    rows = [
        (
            "YOLOv8s",
            metric_text(summary, "YOLOv8s-MultiData"),
            "Not selected for final embedded demo because it is the largest model and was used as the accuracy upper bound.",
            "Best accuracy, but less suitable for the low-power real-time deployment target.",
        ),
        (
            "YOLO11n INT8",
            metric_text(summary, "YOLO11n-MultiData"),
            "Validated on STM32MP257 NPU: 43.20 ms, 23.15 FPS.",
            "Best nano accuracy, but slower than YOLOv8n on the target NPU.",
        ),
        (
            "YOLO26n INT8",
            metric_text(summary, "YOLO26n-MultiData"),
            "Validated on STM32MP257 NPU: 44.08 ms, 22.69 FPS; lower video detection stability.",
            "Compact, but not the strongest accuracy or deployment candidate.",
        ),
        (
            "YOLOv8n INT8",
            metric_text(summary, "YOLOv8n-MultiData"),
            "Validated on STM32MP257 NPU: 35.70 ms, 28.01 FPS; stable sample-video detections.",
            "Selected final model because it provides the best verified NPU deployment trade-off.",
        ),
    ]
    add_table(doc, ["Model", "Accuracy and size evidence", "NPU deployment evidence", "Decision"], rows)
    doc.add_paragraph(
        "Therefore, the parameter/mAP plot supports the accuracy-complexity part of the "
        "decision, while the NPU comparison table provides the deployment feasibility evidence. "
        "Together, they justify selecting YOLOv8n INT8 as the final model for the STM32MP257 "
        "demonstration instead of selecting the highest-mAP model alone."
    )


def add_discussion(doc: Document) -> None:
    doc.add_heading("4. Discussion", level=1)
    doc.add_paragraph(
        "The results support the clinical and engineering feasibility of the project. The "
        "clinical relevance comes from the direct link between model output and procedural "
        "attention: a highlighted suspected polyp can prompt closer inspection during "
        "colonoscopy. The engineering contribution is the measured path from public datasets "
        "to NPU-executed INT8 inference on STM32MP257."
    )
    doc.add_paragraph(
        "The model comparison shows a clear accuracy-deployment trade-off. YOLOv8s is the "
        "best accuracy model, but its 11.14M parameters and higher compute make it less "
        "attractive for low-power real-time deployment. YOLO11n is the strongest nano-scale "
        "accuracy candidate and was successfully converted to a full-integer INT8 .nb model; "
        "however, its measured STM32MP257 NPU latency was 43.20 ms, slower than YOLOv8n. "
        "YOLO26n is compact, but its full-integer INT8 NPU version was also slower and "
        "produced lower detection stability in the sample video. Therefore, YOLOv8n is the "
        "most defensible final demonstration model."
    )
    doc.add_paragraph(
        "The preprocessing experiments are also important. They show that achieving real-time "
        "performance is not only a neural-network problem. On an embedded system, image decode, "
        "resize, channel conversion, quantization, NMS, and overlay rendering must also be "
        "measured. Optimizing preprocessing from a float-based Python route to a direct INT8 "
        "letterbox route improved the usable video pipeline while preserving model geometry."
    )


def add_limitations(doc: Document) -> None:
    doc.add_heading("5. Limitations", level=1)
    for text in [
        "This is a research prototype and not diagnostic or regulatory-grade software.",
        "The datasets are public research datasets and do not replace prospective patient-level clinical validation.",
        "The video demonstration uses dataset-derived video rather than direct integration with a clinical colonoscope.",
        "INT8 quantization and NPU conversion can change score calibration and may require threshold tuning.",
        "YOLOv8n 384 and 320 full-integer INT8 TFLite models were prepared for future speed tests, but their .nb accuracy and NPU latency were not finalized in this report.",
    ]:
        doc.add_paragraph(text, style="List Bullet")


def add_conclusion(doc: Document) -> None:
    doc.add_heading("6. Conclusion", level=1)
    doc.add_paragraph(
        "This project demonstrates a clinically motivated and technically complete edge-AI "
        "prototype for colonoscopy polyp detection. Multi-dataset training and external "
        "validation provide a stronger evaluation than a single-dataset experiment. YOLOv8s "
        "achieves the best external accuracy, but YOLOv8n is selected as the final deployment "
        "model because it has the strongest verified STM32MP257 NPU path: full-integer INT8 "
        "input/output, approximately 35.7 ms NPU inference, and a functional video inference "
        "pipeline. The final result is suitable for a biomedical AI final project because it "
        "connects clinical motivation, model comparison, quantization, embedded deployment, "
        "and measured accuracy-latency trade-offs."
    )


def add_reproducibility(doc: Document) -> None:
    doc.add_heading("7. Reproducibility Package", level=1)
    rows = [
        ("Multi-dataset YAML", "data/processed/polyp-multidata/polyp-multidata-local.yaml"),
        ("Training config", "configs/training_multidata_augmented.yaml"),
        ("Evaluation CSV", "runs/multidata_compare_eval_4models/multidata_eval.csv"),
        ("Model summary CSV", "runs/multidata_compare_eval_4models/model_summary.csv"),
        ("YOLOv8n NPU model", "models/int8/yolov8n_polyp_416_full_integer_int8_1.nb"),
        ("YOLO11n NPU model", "models/int8/yolo11n_polyp_416_full_integer_int8_1.nb"),
        ("YOLO26n NPU model", "models/int8/yolo26n_polyp_416_full_integer_int8_1.nb"),
        ("STM32 NPU comparison CSV", "runs/stm32_inference/stm32_npu_model_comparison.csv"),
        ("STM32 video inference script", "deployment/stai_yolo_video.py"),
    ]
    add_table(doc, ["Item", "Path"], rows)


def add_references(doc: Document) -> None:
    doc.add_heading("References", level=1)
    for ref in REFERENCES:
        doc.add_paragraph(ref, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value


def add_optional_figure(doc: Document, path: Path, caption: str, explanation: str, width: float = 6.2) -> None:
    if not path.exists():
        doc.add_paragraph(f"[Missing figure: {path}]")
        return
    doc.add_picture(str(path), width=Inches(width))
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(explanation)


def metric_text(summary: list[dict[str, str]], model_name: str) -> str:
    row = next(row for row in summary if row["model"] == model_name)
    return (
        f"{float(row['params_m']):.2f}M parameters; "
        f"internal mAP50-95 {float(row['internal_map5095']):.3f}; "
        f"external average mAP50-95 {float(row['external_avg_map5095']):.3f}."
    )


def add_summary_table(doc: Document, summary: list[dict[str, str]]) -> None:
    headers = [
        "Model",
        "Params (M)",
        "GFLOPs",
        "PT MB",
        "Internal mAP50-95",
        "External Avg mAP50-95",
        "External Avg Recall",
        "CVC",
        "ColonDB",
        "ETIS",
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
    for row in summary:
        values = [
            row["model"].replace("-MultiData", ""),
            f"{float(row['params_m']):.2f}",
            f"{float(row['gflops']):.2f}",
            f"{float(row['weight_mb']):.2f}",
            f"{float(row['internal_map5095']):.3f}",
            f"{float(row['external_avg_map5095']):.3f}",
            f"{float(row['external_avg_recall']):.3f}",
            f"{float(row['clinicdb_map5095']):.3f}",
            f"{float(row['colondb_map5095']):.3f}",
            f"{float(row['etis_map5095']):.3f}",
        ]
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = value


def add_detail_table(doc: Document, detail: list[dict[str, str]]) -> None:
    headers = ["Model", "Dataset", "Precision", "Recall", "mAP50", "mAP50-95"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
    for row in detail:
        cells = table.add_row().cells
        values = [
            row["model"].replace("-MultiData", ""),
            row["dataset"],
            f"{float(row['precision']):.3f}",
            f"{float(row['recall']):.3f}",
            f"{float(row['map50']):.3f}",
            f"{float(row['map50_95']):.3f}",
        ]
        for cell, value in zip(cells, values):
            cell.text = value


def make_figures(summary: list[dict[str, str]]) -> list[tuple[Path, str, str]]:
    labels = [row["model"].replace("-MultiData", "") for row in summary]
    internal = [float(row["internal_map5095"]) for row in summary]
    external = [float(row["external_avg_map5095"]) for row in summary]
    params = [float(row["params_m"]) for row in summary]
    npu_labels = ["CPU ONNX", "YOLO26n old .nb", "YOLO11n INT8", "YOLO26n INT8", "YOLOv8n INT8"]
    npu_latency = [879.63, 422.70, 43.20, 44.08, 35.70]
    pre_labels = ["Original", "Fast INT8\nletterbox", "GStreamer\nresize"]
    pre_ms = [51.64, 17.47, 7.60]
    wall_fps = [10.68, 17.44, 21.19]

    paths: list[tuple[Path, str, str]] = []

    fig = FIG_DIR / "accuracy_internal_external.png"
    x = range(len(labels))
    plt.figure(figsize=(8, 4.3))
    plt.bar([i - 0.18 for i in x], internal, width=0.36, label="Internal")
    plt.bar([i + 0.18 for i in x], external, width=0.36, label="External avg")
    plt.xticks(list(x), labels, rotation=15, ha="right")
    plt.ylabel("mAP50-95")
    plt.ylim(0, 0.85)
    plt.legend()
    plt.tight_layout()
    plt.savefig(fig, dpi=220)
    plt.close()
    paths.append((
        fig,
        "Figure 1. Internal holdout and external-average mAP50-95 for the four trained YOLO models.",
        "This figure compares in-distribution performance with external-dataset generalization. "
        "The gap between internal and external mAP50-95 indicates domain shift across public "
        "endoscopy datasets, which is why external testing is more persuasive than reporting "
        "a single internal split alone.",
    ))

    fig = FIG_DIR / "params_vs_external_accuracy.png"
    plt.figure(figsize=(7, 4.3))
    plt.scatter(params, external, s=90)
    for label, p, m in zip(labels, params, external):
        plt.annotate(label, (p, m), xytext=(6, 5), textcoords="offset points")
    plt.xlabel("Parameters (M)")
    plt.ylabel("External average mAP50-95")
    plt.grid(alpha=0.3)
    plt.tight_layout()
    plt.savefig(fig, dpi=220)
    plt.close()
    paths.append((
        fig,
        "Figure 2. Accuracy-deployment trade-off: parameter count versus external-average mAP50-95.",
        "This figure is the first part of the model-selection argument. YOLOv8s is the "
        "accuracy leader but is clearly separated by model size. The nano-scale models occupy "
        "a smaller-parameter region, so the final decision must also consider measured NPU "
        "latency and video stability rather than parameter count or mAP alone.",
    ))

    fig = FIG_DIR / "npu_latency.png"
    plt.figure(figsize=(7.5, 4.3))
    plt.bar(npu_labels, npu_latency)
    plt.ylabel("Latency (ms/frame)")
    plt.yscale("log")
    plt.xticks(rotation=12, ha="right")
    plt.tight_layout()
    plt.savefig(fig, dpi=220)
    plt.close()
    paths.append((
        fig,
        "Figure 3. STM32MP257 inference latency comparison. Log scale highlights the improvement from CPU ONNX to INT8 NPU execution.",
        "The latency comparison demonstrates the practical value of full-integer NPU "
        "deployment. YOLOv8n INT8 achieved the lowest measured NPU latency among the validated "
        ".nb models, which is the main deployment reason it was selected as the final demo model.",
    ))

    fig = FIG_DIR / "pipeline_preprocess_wallfps.png"
    x = range(len(pre_labels))
    fig_obj, ax1 = plt.subplots(figsize=(7.5, 4.3))
    ax1.bar([i - 0.18 for i in x], pre_ms, width=0.36, label="Preprocess ms")
    ax1.set_ylabel("Preprocess latency (ms)")
    ax2 = ax1.twinx()
    ax2.bar([i + 0.18 for i in x], wall_fps, width=0.36, color="tab:orange", label="Wall FPS")
    ax2.set_ylabel("End-to-end wall FPS")
    ax1.set_xticks(list(x))
    ax1.set_xticklabels(pre_labels)
    lines, labels1 = ax1.get_legend_handles_labels()
    lines2, labels2 = ax2.get_legend_handles_labels()
    ax1.legend(lines + lines2, labels1 + labels2, loc="upper center")
    fig_obj.tight_layout()
    fig_obj.savefig(fig, dpi=220)
    plt.close(fig_obj)
    paths.append((
        fig,
        "Figure 4. Preprocessing optimization reduced Python overhead and improved end-to-end video throughput.",
        "This figure shows that the embedded pipeline bottleneck was not limited to neural "
        "network inference. Direct INT8 letterbox preprocessing reduced CPU-side overhead while "
        "preserving YOLO geometry, making it the preferred real-time video route.",
    ))

    return paths


if __name__ == "__main__":
    main()
